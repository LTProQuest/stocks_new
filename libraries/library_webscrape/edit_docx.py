import os
from docx import Document

from docx import Document  
doc = Document()
import re



"""
for line_number, line in enumerate(cv_preview):
    try:
        next_line = cv_preview[line_number+1]
    except:
        pass    
    if ((line.startswith("•") == False) and (next_line.startswith("•") == True)):
        cv_preview[line_number] = line + "caught"

holder = ""
for line in cv_preview:
    holder.join(line)

print("holder - ", holder)    
#print("".join(cv_preview))
"""

def depersonalise_via_regex(text):
    email_regex = re.compile(r"([a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]{0,8})")
    email_regex2 = re.compile(r"[A-Za-z0-9\.\+_-]+@[A-Za-z0-9\._-]+\.[a-zA-Z]*$")
    postcode_regex = re.compile(r'[,a-zA-Z0-9 ]+[A-Za-z]{1,2}[0-9R][0-9A-Za-z]? [0-9][A-Za-z]{2}')
    url_regex = re.compile('http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\(\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+')
    location_regex2 = re.compile(r'[\d]{1,2} +[a-zA-Z]{1,15} +(?:Park|Road|Hill|Lane|London|Green|Avenue|Green|Way)')
    location_regex = re.compile(r"Park,|Road,|Hill,|Lane,|London,|Green,|Avenue,") #in case postcode is not provided which would stop postcode regex clearing address
    phone_regex = re.compile("\+? {0,2}\d+ {0,2}[(-]?\d(?:[ \d]*\d)?[)-]? {0,2}\d+[/ -]?\d+[/ -]?\d+(?: *- *\d+)?")
    regex_list = [email_regex,email_regex2, location_regex2 , postcode_regex, url_regex, phone_regex]
    for regex in regex_list:
        for match in re.findall(regex, text):
            replacement_text = match[0:2] + ((len(match) - 2)*"x")
            
            text = text.replace(match, replacement_text)

    for match in re.findall(location_regex, text):
            replacement_text = ((len(match))*"x")
            
            text = text.replace(match, replacement_text)        
    return text    


def docx_replace_multiple_strings(file_path, old_texts, output_directory):
    doc = Document(file_path)
    address_texts = ["lane,", "road,", "street,", "london", "address:"]
    for section in doc.sections:
        header = section.header
        footer = section.footer
        for paragraph in header.paragraphs:
            paragraph.text = ""
        for paragraph in footer.paragraphs:
            paragraph.text = ""    

    #see the actual text of tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = depersonalise_via_regex(cell.text)
                for text in old_texts:
                    if cell.text.lower().find(text.lower()) != -1:
                        cell.text = ""
                        
                

    for p in doc.paragraphs:
        p.text = depersonalise_via_regex(p.text)
     
        for old_text in old_texts:
            new_text = old_text[0:2] + ((len(old_text) - 2) * "x")
            if p.text.lower().find(old_text.lower()) != -1:
                
                pattern = re.compile(old_text, re.IGNORECASE)
                p.text = pattern.sub(new_text, p.text)
                #p.text = p.text.replace(old_text , new_text)
                
                # style = p.style
                
                # p.style = style
    path, filename = os.path.split(file_path)
    edited_cv_file_path = output_directory  + filename
    doc.save(edited_cv_file_path)
    return edited_cv_file_path




#string = "modo@hotmail.com needs removing ? 29919-depersonalised.pdf – remove address “27 Old Gloucester Street, London” if poss. 29920-depersonalised.pdf – remove LinkedIn address www.linkedin.com/in/courtney-omant"
#print(depersonalise_via_regex(string))
#old_texts = ["Martyn, Braithwate"]
#file_path = r"C:\Users\luket\Desktop\test_space\storage\cv_storage\7000045 - CWJobs).doc"
#output_directory = "storage\edited_cv_storage"
#docx_replace_multiple_strings(file_path, old_texts, output_directory)

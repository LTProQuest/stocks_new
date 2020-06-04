import os
from docx import Document

from docx import Document  
doc = Document()
import re



"""
for line_number, line in enumerate(cv_preview):
    try:
        next_line = cv_preview[line_number+1]
    except:
        pass    
    if ((line.startswith("•") == False) and (next_line.startswith("•") == True)):
        cv_preview[line_number] = line + "caught"

holder = ""
for line in cv_preview:
    holder.join(line)

print("holder - ", holder)    
#print("".join(cv_preview))
"""



def depersonalise_via_regex(text, names):
    email_regex = re.compile(r"([a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]{0,8})")
    email_regex2 = re.compile(r"[A-Za-z0-9\.\+_-]+@[A-Za-z0-9\._-]+\.[a-zA-Z]*$")
    postcode_regex = re.compile(r'[,a-zA-Z0-9 ]{0,40} | [A-Za-z]{1,2}[0-9R][0-9A-Za-z]? [0-9][A-Za-z]{2}')
    url_regex = re.compile('http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\(\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+')
    name_regex_1 = re.compile(names[0])
    name_regex_2 = re.compile(names[1])
    locations_regex = re.compile(r"Park|Road|Hill|Lane|London") #in case postcode is not provided which would stop postcode regex clearing address
    phone_regex = re.compile(r"[+\-\d ]{8,}")
    regex_list = [email_regex,email_regex2,locations_regex , postcode_regex, url_regex, phone_regex, name_regex_1, name_regex_2]
    for regex in regex_list:
        for match in re.findall(regex, text):
            replacement_text = match[0:2] + ((len(match) - 2)*"x")
            
            text = text.replace(match, replacement_text)
    return text    

def docx_replace_multiple_strings(file_path, old_texts, output_directory):
    doc = Document(file_path)
   
    for section in doc.sections:
        header = section.header
        footer = section.footer
        for paragraph in header.paragraphs:
            paragraph.text = ""
        for paragraph in footer.paragraphs:
            paragraph.text = ""    
    for p in doc.paragraphs:
        
        p.text = depersonalise_via_regex(p.text,old_texts)
        for old_text in old_texts:
            new_text = old_text[0:2] + ((len(old_text) - 2) * "x")
            if old_text.lower().find(p.text.lower()) != -1:
                
                pattern = re.compile(old_text, re.IGNORECASE)
                p.text = pattern.sub(new_text, p.text)
                #p.text = p.text.replace(old_text , new_text)
                
                # style = p.style
                
                # p.style = style
    path, filename = os.path.split(file_path)
    edited_cv_file_path = output_directory + "/" + filename
    doc.save(edited_cv_file_path)
    print("new edit doc app not being bypassed")
    return edited_cv_file_path

#header = document.sections[0].header
#header.text = 'foobar'

file_path = (r"C:\Users\luket\Desktop\test_space\storage\cv_storage\17358445 - CWJobs).docx")
output_directory = "storage/edited_cv_storage"
old_texts = ["hi","bye"]
docx_replace_multiple_strings(file_path, old_texts, output_directory)